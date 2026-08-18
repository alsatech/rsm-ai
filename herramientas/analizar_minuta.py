#!/usr/bin/env python3
"""Analizador de minutas — herramienta interna de desarrollo (Alfredo Salas).

NO es un módulo del sistema RSM. Compara el contenido de una minuta operativa
contra el estado real del sistema (CLAUDE.md, PROGRESS.md, CHANGELOG.md, apps
Django) usando la API de Anthropic, y genera un reporte con gaps, mejoras y
prompts listos para pegar en Claude Code.

Uso:
    python herramientas/analizar_minuta.py minutas/2026-08-12.pdf
    python herramientas/analizar_minuta.py minutas/2026-08-12.docx
    python herramientas/analizar_minuta.py minutas/2026-08-12.txt
"""

import os
import re
import sys
from datetime import date
from pathlib import Path

PROJECT_ROOT = Path(__file__).resolve().parent.parent
ENV_PATH = PROJECT_ROOT / "backend" / ".env"
CLAUDE_MD_PATH = PROJECT_ROOT / "CLAUDE.md"
PROGRESS_MD_PATH = PROJECT_ROOT / "PROGRESS.md"
CHANGELOG_MD_PATH = PROJECT_ROOT / "CHANGELOG.md"
BACKEND_APPS_PATH = PROJECT_ROOT / "backend" / "apps"
LOG_PATH = PROJECT_ROOT / "MINUTAS_ANALIZADAS.md"

MODEL = "claude-opus-5"
MAX_CHARS_CONTEXTO = 60000  # por archivo de referencia, para no inflar el prompt


def cargar_env(ruta: Path) -> None:
    """Carga variables desde backend/.env sin depender de Django."""
    if not ruta.exists():
        return
    for linea in ruta.read_text(encoding="utf-8").splitlines():
        linea = linea.strip()
        if not linea or linea.startswith("#") or "=" not in linea:
            continue
        clave, _, valor = linea.partition("=")
        os.environ.setdefault(clave.strip(), valor.strip().strip('"').strip("'"))


def leer_texto_referencia(ruta: Path) -> str:
    if not ruta.exists():
        return "(no disponible)"
    return ruta.read_text(encoding="utf-8", errors="ignore")[:MAX_CHARS_CONTEXTO]


def listar_apps_backend() -> str:
    if not BACKEND_APPS_PATH.exists():
        return "(no se encontró backend/apps/)"
    apps = sorted(
        p.name
        for p in BACKEND_APPS_PATH.iterdir()
        if p.is_dir() and not p.name.startswith("__") and (p / "models.py").exists()
    )
    return ", ".join(apps) if apps else "(sin apps detectadas)"


def extraer_texto_pdf(ruta: Path) -> str:
    from pypdf import PdfReader

    lector = PdfReader(str(ruta))
    return "\n".join(pagina.extract_text() or "" for pagina in lector.pages).strip()


def extraer_texto_docx(ruta: Path) -> str:
    from docx import Document

    documento = Document(str(ruta))
    partes = [p.text for p in documento.paragraphs if p.text.strip()]
    for tabla in documento.tables:
        for fila in tabla.rows:
            celdas = [c.text.strip() for c in fila.cells if c.text.strip()]
            if celdas:
                partes.append(" | ".join(celdas))
    return "\n".join(partes).strip()


def extraer_texto_txt(ruta: Path) -> str:
    try:
        return ruta.read_text(encoding="utf-8").strip()
    except UnicodeDecodeError:
        return ruta.read_text(encoding="latin-1").strip()


def extraer_texto(ruta: Path) -> str:
    extension = ruta.suffix.lower()
    if extension == ".pdf":
        return extraer_texto_pdf(ruta)
    if extension == ".docx":
        return extraer_texto_docx(ruta)
    if extension == ".txt":
        return extraer_texto_txt(ruta)
    raise ValueError(f"Extensión no soportada: {extension} (usar .pdf, .docx o .txt)")


def extraer_fecha(ruta: Path) -> str:
    coincidencia = re.search(r"\d{4}-\d{2}-\d{2}", ruta.stem)
    return coincidencia.group(0) if coincidencia else date.today().isoformat()


def construir_system_prompt(fecha: str) -> str:
    claude_md = leer_texto_referencia(CLAUDE_MD_PATH)
    progress_md = leer_texto_referencia(PROGRESS_MD_PATH)
    changelog_md = leer_texto_referencia(CHANGELOG_MD_PATH)
    apps_backend = listar_apps_backend()

    return f"""Eres un asistente de desarrollo de software especializado en
analizar minutas operativas de la Reserva Santa Margarita AC
y compararlas contra el sistema RSM Sistema que está siendo
desarrollado.

Tu tarea es:
1. Leer la minuta proporcionada
2. Identificar CUALQUIER mención de procesos, problemas,
   necesidades o actividades que:
   a) Ya están cubiertas por un módulo existente del sistema
   b) NO están cubiertas por ningún módulo actual
   c) Están parcialmente cubiertas pero podrían mejorarse
3. Generar propuestas concretas y accionables

El sistema RSM tiene estos módulos (CLAUDE.md):
{claude_md}

Apps Django implementadas actualmente en backend/apps/: {apps_backend}

Estado actual del proyecto (PROGRESS.md):
{progress_md}

Historial de cambios (CHANGELOG.md):
{changelog_md}

Formato de respuesta OBLIGATORIO:

## ANÁLISIS DE MINUTA — {fecha}

### ✅ YA CUBIERTO POR EL SISTEMA
Lista de lo que se menciona en la minuta y ya existe en el sistema.
Por cada item: qué módulo lo cubre y cómo.

### ⚠️ PARCIALMENTE CUBIERTO
Lista de features que existen pero podrían mejorarse basándose
en lo que dice la minuta. Por cada item: qué falta agregar.

### 🆕 NO ESTÁ EN EL SISTEMA — PROPUESTAS NUEVAS
Lista de necesidades detectadas en la minuta que NO tienen
módulo correspondiente. Por cada propuesta:
- Descripción de qué se necesita
- En qué módulo existente encajaría (o si requiere módulo nuevo)
- Prioridad sugerida: Alta / Media / Baja
- Estimado de complejidad: Simple / Medio / Complejo

### 📋 PROMPTS SUGERIDOS PARA CLAUDE CODE
Por cada propuesta nueva o mejora, generar el prompt listo
para pegar en Claude Code. Formato:
---
MÓDULO: [nombre]
CAMBIO: [descripción breve]
PROMPT:
[prompt completo listo para usar]
---

### 🗓️ RESUMEN EJECUTIVO
En máximo 5 líneas: qué es lo más importante que encontraste
en esta minuta que el sistema debería atender pronto.

Soporta minutas en español e inglés — responde siempre en español."""


def llamar_claude(system_prompt: str, texto_minuta: str) -> str:
    import anthropic

    api_key = os.environ.get("ANTHROPIC_API_KEY")
    if not api_key:
        raise RuntimeError(
            f"No se encontró ANTHROPIC_API_KEY. Defínela en {ENV_PATH} "
            "o como variable de entorno."
        )

    client = anthropic.Anthropic(api_key=api_key)
    respuesta = client.messages.create(
        model=MODEL,
        max_tokens=8000,
        system=system_prompt,
        messages=[
            {
                "role": "user",
                "content": f"Aquí está el contenido de la minuta a analizar:\n\n{texto_minuta}",
            }
        ],
    )
    return "".join(bloque.text for bloque in respuesta.content if bloque.type == "text")


def contar_items_seccion(texto: str, encabezado: str) -> int:
    """Cuenta bullets de primer nivel ('- ') dentro de una sección ### dada."""
    inicio = re.search(rf"^###.*{re.escape(encabezado)}", texto, re.MULTILINE)
    if not inicio:
        return 0
    resto = texto[inicio.end():]
    fin = re.search(r"^### ", resto, re.MULTILINE)
    seccion = resto[: fin.start()] if fin else resto
    return len(re.findall(r"^- ", seccion, re.MULTILINE))


def actualizar_log(fecha: str, nombre_archivo: str, propuestas_nuevas: int, mejoras: int) -> None:
    if not LOG_PATH.exists():
        LOG_PATH.write_text(
            "# MINUTAS_ANALIZADAS.md\n"
            "> Log de análisis generados por `herramientas/analizar_minuta.py`.\n\n"
            "| Fecha | Archivo | Propuestas nuevas | Mejoras |\n"
            "|-------|---------|--------------------|---------|\n",
            encoding="utf-8",
        )
    fila = f"| {fecha} | {nombre_archivo} | {propuestas_nuevas} propuestas nuevas | {mejoras} mejoras |\n"
    with LOG_PATH.open("a", encoding="utf-8") as f:
        f.write(fila)


def main() -> None:
    if len(sys.argv) != 2:
        print("Uso: python herramientas/analizar_minuta.py <ruta_minuta>")
        sys.exit(1)

    ruta_minuta = Path(sys.argv[1])
    if not ruta_minuta.is_absolute():
        ruta_minuta = Path.cwd() / ruta_minuta

    if not ruta_minuta.exists():
        print(f"❌ Archivo no encontrado: {ruta_minuta}")
        sys.exit(1)

    cargar_env(ENV_PATH)

    print(f"📄 Extrayendo texto de: {ruta_minuta.name}")
    try:
        texto_minuta = extraer_texto(ruta_minuta)
    except ValueError as error:
        print(f"❌ {error}")
        sys.exit(1)

    if not texto_minuta.strip():
        print("❌ La minuta no tiene contenido legible")
        sys.exit(1)

    fecha = extraer_fecha(ruta_minuta)
    print(f"🧠 Consultando a Claude ({MODEL})...")

    try:
        analisis = llamar_claude(construir_system_prompt(fecha), texto_minuta)
    except Exception as error:
        print(f"❌ Error al llamar a la API de Anthropic: {error}")
        print("   Verifica que ANTHROPIC_API_KEY esté configurada correctamente en backend/.env")
        sys.exit(1)

    print("\n" + "=" * 70)
    print(analisis)
    print("=" * 70 + "\n")

    ruta_salida = ruta_minuta.parent / f"{ruta_minuta.stem}_analisis.md"
    ruta_salida.write_text(analisis, encoding="utf-8")
    print(f"✅ Análisis guardado en: {ruta_salida}")

    propuestas_nuevas = contar_items_seccion(analisis, "NO ESTÁ EN EL SISTEMA")
    mejoras = contar_items_seccion(analisis, "PARCIALMENTE CUBIERTO")
    actualizar_log(fecha, ruta_minuta.name, propuestas_nuevas, mejoras)
    print(f"✅ Registro actualizado en: {LOG_PATH.name}")


if __name__ == "__main__":
    main()
